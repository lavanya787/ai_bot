import os
import re
import json
import fitz  # PyMuPDF
import docx
import PyPDF2
import numpy as np
import logging
import tempfile
import pandas as pd
from PIL import Image
from langdetect import detect, DetectorFactory
from contextlib import contextmanager
from typing import List
from deep_translator import GoogleTranslator
from concurrent.futures import ThreadPoolExecutor
from docx import Document
import pytesseract
import pdfplumber


# ---------------------- Logging Setup ----------------------
log_dir = "logs"
os.makedirs(log_dir, exist_ok=True)
log_file = os.path.join(log_dir, "processor.log")

logger = logging.getLogger("processor")
logger.setLevel(logging.INFO)

# Console Handler
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)

# File Handler (append mode explicitly set)
file_handler = logging.FileHandler(log_file, mode='a', encoding="utf-8")
file_handler.setLevel(logging.INFO)

# Formatter
formatter = logging.Formatter("[%(asctime)s] [%(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
console_handler.setFormatter(formatter)
file_handler.setFormatter(formatter)

# Add handlers only once
if not logger.handlers:
    logger.addHandler(console_handler)
    logger.addHandler(file_handler)

# Reduce log noise from other libraries
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("urllib3").setLevel(logging.WARNING)


# Optional dependencies
try:
    import pdfplumber
    USE_PDFPLUMBER = True
except ImportError:
    USE_PDFPLUMBER = False

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    nlp.max_length = 100_000_000
    SPACY_AVAILABLE = True
except Exception:
    nlp = None
    SPACY_AVAILABLE = False

try:
    import easyocr
    easyocr_reader = easyocr.Reader(['en'], gpu=True)
    EASY_OCR_AVAILABLE = True
except Exception:
    EASY_OCR_AVAILABLE = False

# ---------------------- Utils ----------------------

@contextmanager
def open_tempfile(file, suffix):
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file.read())
        tmp.flush()
        yield tmp.name
    os.remove(tmp.name)

def detect_language(text: str) -> str:
    try:
        lang = detect(text)
        logger.info(f"Detected language: {lang}")
        return lang
    except:
        return "unknown"

def translate_to_english(text: str) -> str:
    try:
        return GoogleTranslator(source='auto', target='en').translate(text)
    except Exception as e:
        logger.warning(f"Translation failed: {e}")
        return text

def preprocess_extracted_text(text: str) -> str:
    text = re.sub(r'\n{2,}', '\n', text)
    text = re.sub(r'\s{2,}', ' ', text)
    text = re.sub(r'\.{4,}', '.', text)
    text = re.sub(r'Page\s+\d+', '', text, flags=re.IGNORECASE)
    return text.strip()

# ------------------ Text Extractors ------------------

def extract_text_from_image(file):
    try:
        if EASY_OCR_AVAILABLE:
            result = easyocr_reader.readtext(np.array(Image.open(file)), detail=0)
            return " ".join(result)
        else:
            return pytesseract.image_to_string(Image.open(file)).strip()
    except Exception as e:
        return f"[OCR ERROR] {e}"

def extract_text_from_txt(file):
    try:
        file.seek(0)
        return file.read().decode("utf-8")
    except Exception as e:
        return f"[TXT ERROR] {e}"

def extract_text_from_csv(file):
    try:
        file.seek(0)
        df = pd.read_csv(file)
        return "\n".join(df.astype(str).apply(" | ".join, axis=1))
    except Exception as e:
        return f"[CSV ERROR] {e}"

def extract_text_from_json(file):
    try:
        file.seek(0)
        data = json.load(file)
        return json.dumps(data, indent=2)
    except Exception as e:
        return f"[JSON ERROR] {e}"

def extract_text_from_docx(file):
    try:
        with open_tempfile(file, ".docx") as path:
            doc = Document(path)
            return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"[DOCX ERROR] {e}"

def extract_text_from_pdf(file):
    def ocr_pdf_images(file):
        try:
            file.seek(0)
            images = []
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    images.append(img)

            ocr_text = ""
            for img in images:
                if EASY_OCR_AVAILABLE:
                    result = easyocr_reader.readtext(np.array(img), detail=0)
                    ocr_text += " ".join(result) + "\n"
                else:
                    ocr_text += pytesseract.image_to_string(img) + "\n"

            return ocr_text.strip()
        except Exception as e:
            logger.error(f"[OCR fallback failed] {e}")
            return "[OCR fallback failed]"

    try:
        file.seek(0)
        if USE_PDFPLUMBER:
            with pdfplumber.open(file) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                if text.strip():
                    return text
                else:
                    logger.warning("[pdfplumber] Extracted text is empty.")
    except Exception as e:
        logger.warning(f"[pdfplumber] failed: {e}")

    try:
        file.seek(0)
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            text = "\n".join(page.get_text() for page in doc)
            if text.strip():
                return text
            else:
                logger.warning("[fitz] Extracted text is empty.")
    except Exception as e:
        logger.warning(f"[fitz] failed: {e}")

    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        text = "\n".join([p.extract_text() or "" for p in reader.pages])
        if text.strip():
            return text
        else:
            logger.warning("[PyPDF2] Extracted text is empty.")
    except Exception as e:
        logger.error(f"[PyPDF2] failed: {e}")

    # 🧠 OCR fallback
    logger.warning("Attempting OCR fallback for scanned PDF...")
    file.seek(0)
    return ocr_pdf_images(file)

def extract_text_from_file(file_path, file_name):
    try:
        file_ext = file_name.lower()

        # CSV Files
        if file_ext.endswith('.csv'):
            df = pd.read_csv(file_path)
            content = df.to_string(index=False)
            return content

        # Excel Files
        elif file_ext.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file_path)
            content = df.to_string(index=False)
            return content

        # PDF Files
        elif file_ext.endswith('.pdf'):
            try:
                reader = PyPDF2.PdfReader(file_path)
                all_pages = [page.extract_text() or "" for page in reader.pages]
                content = "\n".join(all_pages).strip()
                if content:
                    return content
            except Exception as e:
                print(f"[PyPDF2] Failed: {e}. Retrying with pdfplumber...")

            try:
                with pdfplumber.open(file_path) as pdf:
                    all_pages = [page.extract_text() or "" for page in pdf.pages]
                content = "\n".join(all_pages).strip()
                return content
            except Exception as e:
                return f"❌ Could not extract text from PDF: {str(e)}"

        # DOCX Files
        elif file_ext.endswith('.docx'):
            doc = docx.Document(file_path)
            content = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            return content

        # DOC Files (old Word format)
        elif file_ext.endswith('.doc'):
            try:
                content = pytesseract.process(file_path).decode('utf-8')
                return content
            except Exception as e:
                return f"❌ Could not extract text from DOC file: {str(e)}"

        # Plain Text Files
        elif file_ext.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        # Fallback for unknown extensions
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

    except Exception as e:
        return f"❌ Could not read the file: {str(e)}"

def extract_text(file):
    name = file.name.lower()

    if name.endswith((".png", ".jpg", ".jpeg")):
        text = extract_text_from_image(file)
    elif name.endswith(".pdf"):
        text = extract_text_from_pdf(file)
    elif name.endswith(".docx"):
        text = extract_text_from_docx(file)
    elif name.endswith(".txt"):
        text = extract_text_from_txt(file)
    elif name.endswith(".csv"):
        text = extract_text_from_csv(file)
    elif name.endswith(".json"):
        text = extract_text_from_json(file)
    else:
        return "[Unsupported file type: only PDF, DOCX, TXT, CSV, JSON, PNG, JPG]"

    text = preprocess_extracted_text(text)

    if not text.strip() or text.startswith("[PDF extraction failed]"):
        logger.warning(f"[EMPTY TEXT] Extracted text is empty or invalid for: {file.name}")
        return "[ERROR] Empty or invalid content"

    lang = detect_language(text)
    if lang != "en" and lang != "unknown":
        logger.info("Translating document to English...")
        text = translate_to_english(text)

    return text

# ------------------ Parallel Chunking ------------------

def chunk_text(text: str, chunk_size: int = 1200) -> List[str]:
    if not text or len(text.strip()) < 50:
        logger.warning("❌ Cannot chunk: Text is empty or too short.")
        return []

    if text.startswith("[ERROR]") or "Unsupported file type" in text:
        logger.warning("⚠️ Skipping chunking due to invalid file type or extraction failure.")
        return []

    try:
        if SPACY_AVAILABLE and nlp:
            doc = nlp(text)
            sentences = [sent.text.strip() for sent in doc.sents]
        else:
            sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    except Exception as e:
        logger.warning(f"Chunk fallback due to error: {e}")
        sentences = text.split(". ")

    chunks = []

    def process_chunk(start_idx):
        chunk = ""
        idx = start_idx
        while idx < len(sentences) and len(chunk) + len(sentences[idx]) < chunk_size:
            chunk += sentences[idx] + " "
            idx += 1
        return chunk.strip()

    with ThreadPoolExecutor(max_workers=4) as executor:
        i = 0
        futures = []
        while i < len(sentences):
            futures.append(executor.submit(process_chunk, i))
            j = i
            chunk_len = 0
            while j < len(sentences) and chunk_len + len(sentences[j]) < chunk_size:
                chunk_len += len(sentences[j])
                j += 1
            i = j

        for future in futures:
            chunk = future.result()
            if len(chunk.split()) >= 10:
                chunks.append(chunk)

    return chunks

# ------------------ Data Extraction for QA Models ------------------

def extract_rows_from_file(file_path):
    ext = file_path.split('.')[-1].lower()

    if ext == 'csv':
        df = pd.read_csv(file_path)
    elif ext in ['xls', 'xlsx']:
        df = pd.read_excel(file_path)
    elif ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        df = pd.DataFrame({'question': lines, 'context': lines, 'is_answer': 1})
    elif ext == 'pdf':
        reader = PyPDF2.PdfReader(file_path)
        text = "\n".join([page.extract_text() or "" for page in reader.pages])
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        df = pd.DataFrame({'question': lines, 'context': lines, 'is_answer': 1})
    elif ext == 'docx':
        doc = docx.Document(file_path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        df = pd.DataFrame({'question': lines, 'context': lines, 'is_answer': 1})
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    if all(col in df.columns for col in ['question', 'context', 'is_answer']):
        return df
    else:
        df = pd.DataFrame({'question': df.iloc[:, 0], 'context': df.iloc[:, 0], 'is_answer': 1})
        return df
